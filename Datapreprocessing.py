from docx import Document
from fpdf import FPDF
import os

# Create Word document
doc = Document()
doc.add_heading('창업사업화 지원사업 사업계획서', 0)

# Add sections to Word document
sections = [
    ("1. 신청현황", [
        "신청 주관기관명: OOO",
        "과제번호: OOOOOOOO",
        "신청 분야: 일반분야",
        "사업 분야: 지식서비스",
        "기술 분야: 정보·통신(SW)",
        "사업비 구성계획 (정부지원금): 100백만원"
    ]),
    ("2. 일반현황", [
        "창업아이템명: 5초 마케팅",
        "산출물: 5초 마케팅 웹사이트 (1개)",
        "창업팀 구성 현황:",
        "- 세일즈매니저: 국내외영업 (경력: OO스테이션 영업팀장 6년)",
        "- 선임개발자: 웹/앱개발 (경력: OO플래닛 개발팀장 6년)",
        "- 경영지원팀장: 경영지원 (경력: OO공사 경영지원팀 4년)",
        "- 선임디자이너: 웹/앱디자인 (예정: 2023.5)"
    ]),
    ("3. 창업아이템 개요", [
        "명칭: 5초 마케팅",
        "범주: 마케팅 플랫폼 (웹)",
        "개요: 373만 소상공인을 위한 AI 마케팅 자동화 솔루션",
        "진출 목표시장: 소상공인 373만 중 병원, 변호사, 헬스클럽, 카페, 미용실 등 12만명",
        "차별성:",
        "- 무료 사용",
        "- 온라인/오프라인 채널 활용",
        "- 맞춤형 마케팅 전략 제공"
    ]),
    ("4. 문제인식 (Problem)", [
        "개발 동기 및 추진 경과:",
        "- 마케팅이 어려운 소상공인 대상 문제 해결을 위한 기획",
        "- 50건 이상의 마케팅 컨설팅 경험을 바탕으로 한 문제 인식",
        "- 소상공인들의 마케팅 활동의 어려움과 무분별한 광고비 지출 문제 인식",
        "개발 목적:",
        "- 소상공인들에게 효율적이고 저렴한 마케팅 전략 제공",
        "- 불필요한 광고비 지출 감소 및 마케팅 성공 경험 선물",
        "목표시장 분석:",
        "- 1차 타깃: 병원, 변호사, 헬스클럽, 카페, 미용실, 온라인 쇼핑몰 사업자 (12만명)",
        "- 경쟁이 치열하고 지불 능력이 높은 시장"
    ]),
    ("5. 실현 가능성 (Solution)", [
        "개발 방안 및 진행 정도:",
        "- 웹사이트 와이어프레임 구성 및 시각화 디자인 완료",
        "- 최적화 업체 매칭 기능 개발 중 (30% 완료)",
        "기술 보호 계획: 기술임치제도 활용, 중소기업 기술보호 정책보험 가입",
        "차별화 방안:",
        "- 맞춤형 마케팅 전략 제공 및 최적화된 업체 매칭",
        "- 온라인/오프라인 마케팅 채널 모두 활용",
        "- 소상공인 부담 최소화를 위한 무료 플랫폼 제공"
    ]),
    ("6. 성장 전략 (Scale-up)", [
        "사업화 방안 (비즈니스 모델):",
        "- 소상공인과 실행사의 매칭 중개 플랫폼",
        "- 사용자와 광고주 간 중개 수수료 20% 기반",
        "매출 예상: 1,500명 유치 시 20억원 예상",
        "목표시장 진출 방안:",
        "- B2B 기관(소상공인협회, 프랜차이즈 협회 등) 공략",
        "- SNS를 통한 무료 마케팅 컨설팅 제공 및 타깃 광고"
    ]),
    ("7. 사업 추진 일정", [
        "서비스 개발: 2022.07",
        "서비스 고도화: 2022 하반기",
        "정식 론칭: 2022.12",
        "MOU 체결: 2023 상반기",
        "해외시장 진출: 2025 상반기 (태국)"
    ])
]

# Add content to Word document
for heading, content in sections:
    doc.add_heading(heading, level=1)
    for paragraph in content:
        doc.add_paragraph(paragraph)

# Save Word document
word_file_path = r'C:\Users\jjj11\OneDrive\바탕 화면\창업사업화_지원사업_사업계획서.docx'
doc.save(word_file_path)

# Create PDF document
pdf = FPDF()
pdf.add_page()

# 폰트 파일 경로 수정
font_path = r'C:\Users\jjj11\OneDrive\바탕 화면\NanumGothic.ttf'  # 나눔고딕 폰트 경로
if not os.path.exists(font_path):
    raise FileNotFoundError(f"Font file not found at: {font_path}")

pdf.add_font('NanumGothic', '', font_path, uni=True)
pdf.set_font('NanumGothic', '', 12)

# Add content to PDF document
pdf.set_font('NanumGothic', '', 16)
pdf.cell(0, 10, "창업사업화 지원사업 사업계획서", ln=True, align="C")
pdf.ln(10)

pdf.set_font('NanumGothic', '', 12)
for heading, content in sections:
    pdf.set_font('NanumGothic', '', 14)
    pdf.cell(0, 10, heading, ln=True)
    pdf.set_font('NanumGothic', '', 12)
    for paragraph in content:
        pdf.multi_cell(0, 10, paragraph)
    pdf.ln(5)

# Save PDF document
pdf_file_path = r'C:\Users\jjj11\OneDrive\바탕 화면\창업사업화_지원사업_사업계획서.pdf'
pdf.output(pdf_file_path)

print(f"Word 파일 저장 경로: {word_file_path}")
print(f"PDF 파일 저장 경로: {pdf_file_path}")


